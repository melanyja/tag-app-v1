from docx import Document
from io import BytesIO
import re
import requests
from docx.shared import Inches
from b2sdk.v2 import B2Api, InMemoryAccountInfo
from logger import logger 
from config import settings 

info = InMemoryAccountInfo()
b2_api = B2Api(info)
b2_api.authorize_account("production", settings.B2_KEY_ID, settings.B2_APP_KEY)
bucket = b2_api.get_bucket_by_name(settings.BUCKET_NAME)

def replace_variables_in_doc(doc: Document, variable_dict: dict):
    def is_image_key(value: str) :
        logger.info("Image found")
        return isinstance(value, str) and value.startswith("tags/") and value.lower().endswith((".png", ".jpg", ".jpeg", ".webp", ".bmp"))

    def insert_image_run(paragraph, b2_key):
        try:
            b2_file = bucket.download_file_by_name(b2_key)
            image_stream = BytesIO()
            b2_file.save(image_stream)
            image_stream.seek(0)

            run = paragraph.add_run()
            run.add_picture(image_stream, width=Inches(2))
        except Exception as e:
            logger.error(f"Failed to insert image from B2: {e}")
 

    def replace_text_or_image(paragraph):
        for key, value in variable_dict.items():
            pattern = r'\{\s*' + re.escape(key) + r'\s*\}'
            matches = list(re.finditer(pattern, paragraph.text))
            if matches:
                if is_image_key(value):
                    paragraph.text = re.sub(pattern, '', paragraph.text)
                    insert_image_run(paragraph, value)
                else:
                    paragraph.text = re.sub(pattern, value, paragraph.text)

    for paragraph in doc.paragraphs:
        replace_text_or_image(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_or_image(paragraph)